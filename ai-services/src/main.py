from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import os
from dotenv import load_dotenv, find_dotenv
import PyPDF2
import io
import random
from urllib.parse import quote_plus
import nltk
from nltk.corpus import wordnet
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.tag import pos_tag
from collections import Counter
import re
import gc
import tempfile
import shutil
import json
from groq import Groq

# Configure Groq
load_dotenv()
GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
groq_available = False
groq_client = None
if GROQ_API_KEY:
    try:
        groq_client = Groq(api_key=GROQ_API_KEY)
        groq_available = True
        print("Groq API successfully configured!")
    except Exception as e:
        print(f"Error configuring Groq: {e}")

def get_groq_completion(prompt: str, model_name="llama-3.3-70b-versatile", response_format=None):
    if not groq_client:
        raise Exception("Groq client not initialized")
        
    completion_args = {
        "model": model_name,
        "messages": [{"role": "user", "content": prompt}],
    }
    
    if response_format:
        completion_args["response_format"] = response_format
        
    try:
        completion = groq_client.chat.completions.create(**completion_args)
        return completion.choices[0].message.content
    except Exception as e:
        print(f"Failed to load primary model {model_name}: {e}. Trying fallback chain...")
        fallback_model = "llama-3.1-8b-instant"
        completion_args["model"] = fallback_model
        completion = groq_client.chat.completions.create(**completion_args)
        return completion.choices[0].message.content

def strip_json_markdown(text: str) -> str:
    """
    Strips markdown code block wrappers (like ```json ... ```) from JSON response strings.
    """
    text = text.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].startswith("```"):
            lines = lines[:-1]
        text = "\n".join(lines).strip()
    
    start_brace = text.find("{")
    start_bracket = text.find("[")
    
    start_idx = -1
    end_idx = -1
    
    if start_brace != -1 and (start_bracket == -1 or start_brace < start_bracket):
        start_idx = start_brace
        end_idx = text.rfind("}")
    elif start_bracket != -1:
        start_idx = start_bracket
        end_idx = text.rfind("]")
        
    if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
        text = text[start_idx:end_idx+1]
        
    return text

def get_relevant_excerpt(material_text: str, question: str, correct_answer: str, user_answer: str, context: str, max_sentences: int = 6, max_chars: int = 2500) -> str:
    """
    Pick the most relevant excerpt from the full study material so the explanation
    model can reason over the actual source instead of only the question snippet.
    """
    if not material_text or not material_text.strip():
        return ""

    clean_text = re.sub(r'\s+', ' ', material_text).strip()
    sentences = [s.strip() for s in sent_tokenize(clean_text) if len(s.strip()) > 20]
    if not sentences:
        return clean_text[:max_chars]

    keyword_sources = " ".join([question or "", correct_answer or "", user_answer or "", context or ""]).lower()
    keywords = {
        token for token in re.findall(r"[a-zA-Z][a-zA-Z\-']+", keyword_sources)
        if len(token) > 3
    }

    scored_sentences = []
    for sentence in sentences:
        sent_lower = sentence.lower()
        score = 0
        for keyword in keywords:
            if keyword in sent_lower:
                score += 2
        if correct_answer and correct_answer.lower() in sent_lower:
            score += 6
        if context and context[:120].lower() in sent_lower:
            score += 4
        if question and question[:120].lower() in sent_lower:
            score += 3
        scored_sentences.append((score, sentence))

    top_sentences = [sentence for score, sentence in sorted(scored_sentences, key=lambda item: item[0], reverse=True) if score > 0][:max_sentences]
    if not top_sentences:
        top_sentences = sentences[:max_sentences]

    excerpt = " ".join(top_sentences)
    return excerpt[:max_chars]

def build_deterministic_explanation(question: str, user_answer: str, correct_answer: str, context: str, source_reference: str, topic: str) -> dict:
    clean_question = re.sub(r'\s+', ' ', (question or '')).strip()
    completed_question = clean_question.replace('________', correct_answer).replace('______', correct_answer).replace('_____', correct_answer)
    completed_question = completed_question.replace('____', correct_answer)

    if not completed_question:
        completed_question = f"The correct phrase is '{correct_answer}' in the given context."

    if source_reference:
        source_excerpt = source_reference[:1200]
    else:
        source_excerpt = context[:1200] if context else completed_question[:1200]

    completed_phrase = completed_question if completed_question else correct_answer
    if completed_phrase == correct_answer and ' and ' in clean_question.lower():
        completed_phrase = f"{correct_answer} and the rest of the technical phrase from the sentence"

    if 'communications technologies' in clean_question.lower() and 'information' in correct_answer.lower():
        real_meaning = (
            "The completed phrase is 'Information and Communications Technologies' (ICT). "
            "It refers to the infrastructure, devices, software, networks, and systems used to process, store, and transmit digital information."
        )
        concept_breakdown = (
            "The blank is not just any noun; it is the word that completes a fixed technical phrase. "
            "Here, 'Information' works because the phrase becomes the standard computing term ICT."
        )
        why_wrong = (
            f"'{user_answer}' does not complete the technical phrase correctly and does not match the definition of modern computing infrastructure."
        )
        memory_tip = "If a blank completes a well-known technical phrase, check the standard term first before guessing a random noun."
    else:
        real_meaning = (
            f"'{correct_answer}' is the exact word or phrase that fits the definition, sentence logic, and subject being discussed in the material."
        )
        concept_breakdown = (
            f"Think of '{correct_answer}' as the word that makes the sentence technically and conceptually correct, not just a related term."
        )
        why_wrong = (
            f"'{user_answer}' sounds related, but it does not fit the sentence structure or the meaning required by the material."
        )
        memory_tip = "Focus on the exact sentence meaning and the technical phrase, not just a familiar-sounding option."

    explanation = (
        f"The completed statement is:\n{completed_phrase}\n\n"
        f"Why '{user_answer}' is wrong:\n{why_wrong}\n\n"
        f"Real meaning in context:\n{real_meaning}\n\n"
        f"Concept breakdown:\n{concept_breakdown}\n\n"
        f"Course logic:\n{source_excerpt}"
    )

    return {
        "explanation": explanation,
        "why_wrong": why_wrong,
        "real_meaning": real_meaning,
        "concept_breakdown": concept_breakdown,
        "memory_tip": memory_tip,
        "suggestions": f"Review the section of the material that defines {correct_answer} and compare it with the surrounding sentence.",
        "source_excerpt": source_excerpt,
        "search_terms": [correct_answer, topic, clean_question[:60]]
    }

def looks_generic_explanation(payload: dict) -> bool:
    text = " ".join([
        str(payload.get("explanation", "")),
        str(payload.get("real_meaning", "")),
        str(payload.get("concept_breakdown", "")),
        str(payload.get("suggestions", ""))
    ]).lower()
    generic_markers = [
        "key term",
        "essential concept",
        "important concept",
        "understanding this concept is essential",
        "reviewing",
        "will help you master this specific concept"
    ]
    return (
        len(str(payload.get("explanation", "")).strip()) < 120
        or any(marker in text for marker in generic_markers)
        or not payload.get("concept_breakdown")
        or not payload.get("real_meaning")
    )


app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- NLP Objective Question Generator ---

def preprocess_text(text):
    """Clean and tokenize text into sentences"""
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    sentences = sent_tokenize(text)
    return [s for s in sentences if len(s.strip()) > 40]

def is_high_quality(sentence):
    """Filter for sentences that sound like facts or definitions"""
    indicators = [
        ' is ', ' are ', ' refers to ', ' defined as ', ' known as ', 
        ' because ', ' therefore ', ' significant ', ' primary ', 
        ' essentially ', ' consists of ', ' includes ', ' process of '
    ]
    # Length check + indicator check
    if len(sentence.split()) < 10:
        return False
        
    return any(ind in sentence.lower() for ind in indicators)

def extract_keywords(sentence):
    """Extract candidate keywords using POS tagging, prioritizing longer nouns"""
    tokens = word_tokenize(sentence)
    tagged = pos_tag(tokens)
    
    # Priority 1: Proper Nouns (NNP) - usually names of theories, places, things
    # Priority 2: Standard Nouns (NN)
    # Priority 3: Adjectives over 6 chars
    candidates = []
    for word, pos in tagged:
        if len(word) < 4 or not word.isalnum():
            continue
            
        if pos == 'NNP':
            candidates.append((word, pos, 3)) # High priority
        elif pos in ('NN', 'NNS'):
            candidates.append((word, pos, 2))
        elif pos == 'JJ' and len(word) > 6:
            candidates.append((word, pos, 1))
            
    # Sort by priority then by length (longer = more specific)
    candidates.sort(key=lambda x: (x[2], len(x[0])), reverse=True)
    return [(c[0], c[1]) for c in candidates]

def get_wordnet_pos(treebank_tag):
    """Map treebank POS tags to WordNet POS tags"""
    if treebank_tag.startswith('J'):
        return wordnet.ADJ
    elif treebank_tag.startswith('V'):
        return wordnet.VERB
    elif treebank_tag.startswith('N'):
        return wordnet.NOUN
    elif treebank_tag.startswith('R'):
        return wordnet.ADV
    else:
        return wordnet.NOUN

_distractor_cache = {}

def generate_distractors(target_word, pos_tag):
    """Generate distractors using WordNet strategies (hypernyms, hyponyms, coordinates)"""
    cache_key = f"{target_word.lower()}_{pos_tag}"
    if cache_key in _distractor_cache:
        return _distractor_cache[cache_key]
        
    distractors = set()
    wn_pos = get_wordnet_pos(pos_tag)
    
    synsets = wordnet.synsets(target_word, pos=wn_pos)
    if not synsets:
        # Try without POS restriction if nothing found
        synsets = wordnet.synsets(target_word)
        
    if synsets:
        synset = synsets[0]
        
        # 1. Hypernym Strategy (Coordinate terms)
        for hypernym in synset.hypernyms():
            for hyponym in hypernym.hyponyms():
                name = hyponym.lemmas()[0].name().replace('_', ' ')
                if name.lower() != target_word.lower():
                    distractors.add(name)
                    if len(distractors) >= 10: break
            if len(distractors) >= 10: break
            
        # 2. Hyponym Strategy
        if len(distractors) < 5:
            for hyponym in synset.hyponyms():
                name = hyponym.lemmas()[0].name().replace('_', ' ')
                if name.lower() != target_word.lower():
                    distractors.add(name)
                    if len(distractors) >= 15: break

    result = list(distractors)
    _distractor_cache[cache_key] = result
    # Keep cache small but reasonable
    if len(_distractor_cache) > 1000:
        _distractor_cache.clear()
        
    return result

def filter_distractors(distractors, target_word, sentence, pos_tag):
    """Apply filtering rules for high-quality distractors"""
    filtered = []
    target_len = len(target_word)
    
    # Pre-tokenize sentence for POS consistency check
    # (Simple check: distractors should likely be valid in context)
    
    for d in distractors:
        d = d.lower()
        # Rule: Not a synonym or duplicate
        if d == target_word.lower() or d in filtered:
            continue
            
        # Rule: Not a substring of the answer
        if d in target_word.lower() or target_word.lower() in d:
            continue
            
        # Rule: Similar POS (approximated by WordNet match or word structure)
        # Rule: Reasonable length (not wildly different from target)
        if abs(len(d) - target_len) > 8:
            continue
            
        filtered.append(d)
        if len(filtered) >= 3:
            break
            
    # Fallback to similar words if WordNet fails
    fallbacks = ["process", "system", "theory", "method", "concept", "approach", "framework"]
    while len(filtered) < 3:
        f = random.choice(fallbacks)
        if f not in filtered and f != target_word.lower():
            filtered.append(f)
            
    return filtered

def score_question(question_text, answer, distractors):
    """Simple scoring module (0-100)"""
    score = 100
    
    # 1. POS consistency (simple check: if answer is capitalized, distractors should be too - ignored for now)
    
    # 2. Option length similarity
    avg_len = sum(len(o) for o in distractors + [answer]) / 4
    for o in distractors + [answer]:
        if abs(len(o) - avg_len) > 10:
            score -= 10
            
    # 3. Duplicate check
    if len(set([answer] + distractors)) < 4:
        score -= 50
        
    return score

def generate_mcq(sentence, target_keyword_info):
    """Build a complete MCQ object following the pipeline"""
    word, pos = target_keyword_info
    
    # 1. Distractors
    raw_distractors = generate_distractors(word, pos)
    filtered = filter_distractors(raw_distractors, word, sentence, pos)
    
    # 2. Scoring
    q_score = score_question(sentence, word, filtered)
    if q_score < 70:
        return None # Reject low quality
    
    # 3. Format
    options = filtered + [word]
    random.shuffle(options)
    
    # Assign A, B, C, D
    mapping = {0: 'A', 1: 'B', 2: 'C', 3: 'D'}
    answer_letter = ''
    formatted_options = []
    for i, opt in enumerate(options):
        letter = mapping[i]
        formatted_options.append(f"{letter}. {opt}")
        if opt == word:
            answer_letter = letter
            
    return {
        "type": "mcq",
        "question": sentence.replace(word, "________"),
        "options": options,
        "formatted_options": formatted_options,
        "answer": word,
        "correct_letter": answer_letter,
        "explanation": f"The term '{word}' correctly completes the context of this statement found in your study material.",
        "score": q_score
    }

def generate_essay(sentence):
    """Generate an essay/short answer question from a sentence"""
    keywords = extract_keywords(sentence)
    if not keywords:
        return None
    
    # List of generic words to NEVER use as essay topics
    blacklist = {
        'term', 'context', 'material', 'section', 'point', 'study', 'question', 
        'text', 'sentence', 'subject', 'relationship', 'importance', 'overall',
        'significance', 'discussion', 'concept', 'structure', 'element', 'core'
    }
    
    # 1. Focus on NOUNS (NN, NNP, NNS)
    # 2. Filter out blacklist words
    # 3. Require minimum length of 6 for better quality
    candidates = [word for word, pos in keywords 
                 if pos.startswith('N') 
                 and word.lower() not in blacklist 
                 and len(word) >= 6]
    
    if not candidates:
        # Fallback to any noun over 5 chars if no specific candidate found
        candidates = [word for word, pos in keywords if pos.startswith('N') and len(word) > 5]
        
    if not candidates:
        return None
    
    # Pick the LONGEST noun - it is usually the most specific technical term
    subject = max(candidates, key=len)
    
    prompts = [
        f"Explain the significance of '{subject}' based on your study material.",
        f"How does '{subject}' relate to the key themes discussed in this section?",
        f"Discuss the role and impact of '{subject}' as described in the provided text.",
        f"Looking at your material, summarize the most important points regarding '{subject}'."
    ]
    
    return {
        "type": "essay",
        "question": random.choice(prompts),
        "context": sentence,
        "answer": sentence, # Context serves as the reference answer
        "explanation": f"Analyze your response above. Does it correctly address the role of '{subject}'? The reference material says: \"{sentence}\""
    }

# --- API Endpoints ---

@app.get("/")
async def root():
    return {"status": "OK", "service": "TutorBuddy AI Core", "version": "2.0 (NLP Pipeline)"}

@app.get("/health")
async def health():
    return {"status": "OK"}

@app.post("/extract-text")
async def extract_text(file: UploadFile = File(...)):
    filename = file.filename.lower()
    print(f"--- Ultra-Lean Extraction: {filename} ---")
    
    # Use a temporary file to keep RAM free
    with tempfile.NamedTemporaryFile(delete=False) as tmp:
        try:
            shutil.copyfileobj(file.file, tmp)
            tmp_path = tmp.name
            tmp.close() # Close so other libs can open it

            if filename.endswith('.pdf'):
                reader = PyPDF2.PdfReader(tmp_path)
                pages_content = []
                # Extreme safety: 30 pages / 30k chars
                page_limit = min(len(reader.pages), 35)
                
                total_chars = 0
                for i in range(page_limit):
                    page_text = reader.pages[i].extract_text() or ""
                    pages_content.append(page_text)
                    total_chars += len(page_text)
                    if total_chars > 30000: break
                
                return {"text": "".join(pages_content)[:30000]}
                
            elif filename.endswith('.docx'):
                import docx
                doc = docx.Document(tmp_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                return {"text": text[:30000]}
                
            elif filename.endswith('.txt'):
                with open(tmp_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return {"text": f.read(30000)}
            else:
                raise HTTPException(status_code=400, detail="Unsupported format")
        
        except Exception as e:
            print(f"Lean Extraction Error: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Failed: {str(e)}")
        finally:
            # Cleanup!
            if 'tmp_path' in locals() and os.path.exists(tmp_path):
                os.remove(tmp_path)
            # Force RAM cleanup
            gc.collect()
            await file.close()
 
@app.post("/generate-quiz")
async def generate_quiz(data: dict):
    gc.collect() # Clean up before starting
    try:
        text = data.get("text", "")
        num_requested = data.get("num_questions", 35) # Default total 35
        num_mcq = data.get("num_mcq", 30)
        num_essay = data.get("num_essay", 5)
        difficulty = str(data.get("difficulty", "medium")).lower()
        performance_score = data.get("performance_score")
        adaptive_guidance = data.get("adaptive_guidance", "")
        topic_focus = data.get("topicFocus", "")
        custom_template = data.get("custom_template", None)
        
        print(f"--- Quiz Generation Request Received ---")
        print(f"Text length: {len(text)} characters. Targets: {num_mcq} MCQ, {num_essay} Essay. Topic Focus: {topic_focus or 'None'}")
        if custom_template:
            print(f"Active custom template received and will be injected.")
        
        if not text:
            raise HTTPException(status_code=400, detail="No text provided")
 
        # Lowered limit for better stability on free tier
        if len(text) > 30000:
            print(f"Truncating text from {len(text)} to 30,000 chars for performance.")
            text = text[:30000]

        difficulty_guidance = {
            "easy": "revision mode: create simpler questions, clearer wording, more direct recall, and foundational reinforcement.",
            "medium": "balanced mode: mix recall with moderate conceptual reasoning.",
            "hard": "challenge mode: create deeper analytical, scenario-based, and application-heavy questions with subtle distractors."
        }
        selected_guidance = difficulty_guidance.get(difficulty, difficulty_guidance["medium"])
        if adaptive_guidance:
            selected_guidance = f"{selected_guidance} Additional guidance: {adaptive_guidance}"

        topic_instruction = f"IMPORTANT FOCUS: Ensure all generated questions strictly relate to or cover the topic of '{topic_focus}'. If the text mentions '{topic_focus}', aggressively target that area. " if topic_focus else ""

        performance_note = f"The learner's recent performance score is {performance_score}%. " if performance_score is not None else "No recent performance score is available. "

        # --- Premium Google Gemini Path ---
        if groq_available:
            print("Running Premium Groq Quiz Generator...")
            try:
                
                # If an admin template is provided, wrap it securely to prevent JSON breakage.
                template_instruction = f"""
                ### CUSTOM ADMIN INSTRUCTION / TEMPLATE ###
                {custom_template}
                ###########################################
                (Note: You MUST obey the above instructions for the style and tone of the questions, but you MUST STILL output strictly valid JSON matching the schema below.)
                """ if custom_template else ""

                prompt = f"""
                Generate a study quiz based on the text below. 
                The quiz MUST contain exactly {num_mcq} Multiple Choice Questions (MCQ) and exactly {num_essay} Essay/Short Answer questions.
                Adaptive difficulty mode: {difficulty.upper()}.
                {performance_note}
                Use this teaching style: {selected_guidance}
                {topic_instruction}
                
                {template_instruction}
                
                For MCQs:
                - The questions must be deep, conceptual, and check real-world logic or application.
                - Avoid simple blank fill-ins or direct word matching.
                - Match the question style to the adaptive mode above.
                - If difficulty is EASY, use more direct revision prompts and simpler distractors.
                - If difficulty is MEDIUM, balance recall with some reasoning.
                - If difficulty is HARD, ask more demanding, inferential, and application-based questions.
                - Each MCQ must have exactly 4 options.
                - The "options" list must contain the plain texts.
                - The "formatted_options" list must be prefixed with "A. ", "B. ", "C. ", "D. ".
                - One option must be the correct answer.
                - "correct_letter" must be 'A', 'B', 'C', or 'D', matching the correct option.
                - Include a detailed educational "explanation" for why it is correct.
                - A quality "score" between 80 and 100.
                
                For Essays:
                - The questions must ask the student to explain core concepts, relationships between mechanisms, or summary of main themes in the text.
                - Match the difficulty mode: EASY should be revision-focused and guided, HARD should be analytical and evaluative.
                - "context" must be a supporting reference sentence from the text.
                - "answer" must be a comprehensive reference master answer.
                - "explanation" must be educational criteria for a correct response.
                
                Output the result STRICTLY as a JSON object matching this schema:
                {{
                  "questions": [
                     {{
                       "type": "mcq",
                       "question": "question text",
                       "options": ["option 1", "option 2", "option 3", "option 4"],
                       "formatted_options": ["A. option 1", "B. option 2", "C. option 3", "D. option 4"],
                       "answer": "option text matching the correct answer",
                       "correct_letter": "A",
                       "explanation": "detailed educational explanation of why it is correct",
                       "score": 95
                     }},
                     ...
                     {{
                       "type": "essay",
                       "question": "conceptual question about core themes",
                       "context": "supporting context from the text",
                       "answer": "comprehensive reference master answer",
                       "explanation": "educational criteria for a correct response"
                     }}
                  ]
                }}
                
                CRITICAL INSTRUCTION: DO NOT summarize the text. DO NOT write any introductory or concluding text. 
                You must ONLY output the raw JSON object containing the questions array. No markdown, no explanations outside the JSON.
                
                Source material text:
                {text}
                """
                
                text_content = get_groq_completion(prompt, response_format={"type": "json_object"})
                text_content = strip_json_markdown(text_content)
                quiz_data = json.loads(text_content)
                if "questions" in quiz_data and len(quiz_data["questions"]) > 0:
                    print(f"Groq Quiz Generation successful! Generated {len(quiz_data['questions'])} questions.")
                    return {"quiz": quiz_data}
                else:
                    print("Groq response did not contain questions. Falling back to local NLTK...")
            except Exception as e:
                print(f"Groq quiz generation failed ({e}). Falling back to local NLTK...")

        # --- Local NLTK Fallback Path ---
        print("Running Local NLTK Quiz Generator...")
        sentences = preprocess_text(text)
        if not sentences:
            raise HTTPException(status_code=400, detail="Text too short or invalid to generate questions")
            
        random.shuffle(sentences)
        
        quiz_questions = []
        used_words = set()
        
        # Phase 1: MCQs
        mcq_count = 0
        for i, sent in enumerate(sentences):
            if mcq_count >= num_mcq:
                break
                
            # NEW: prioritize "Educational" sentences
            if not is_high_quality(sent) and len(sentences) > num_mcq * 3:
                continue

            if i % 10 == 0:
                print(f"Processing MCQ sentence {i}/{len(sentences)}... Found {mcq_count} so far.")

            keywords = extract_keywords(sent)
            if not keywords or len(keywords) < 2: 
                continue
                
            # Pick a lucky keyword
            random.shuffle(keywords)
            for kw_info in keywords:
                word = kw_info[0]
                if word.lower() in used_words: 
                    continue
                    
                mcq = generate_mcq(sent, kw_info)
                if mcq:
                    quiz_questions.append(mcq)
                    used_words.add(word.lower())
                    mcq_count += 1
                    break
        
        # Phase 1.5: MCQ Fallback (if we still don't have enough MCQs)
        current_mcq = len([q for q in quiz_questions if q['type'] == 'mcq'])
        if current_mcq < num_mcq:
            print(f"MCQ Fallback: only found {current_mcq}/{num_mcq}")
            for sent in sentences:
                if len([q for q in quiz_questions if q['type'] == 'mcq']) >= num_mcq: break
                # Simplified MCQ for fallback: just pick a word
                words = [w for w in word_tokenize(sent) if len(w) > 5 and w.isalnum()]
                for w in words:
                    if w.lower() in used_words: continue
                    mcq = generate_mcq(sent, (w, 'NN'))
                    if mcq:
                        quiz_questions.append(mcq)
                        used_words.add(w.lower())
                        break
        
        # Phase 2: Essays
        essay_candidates = [s for s in sentences if len(s.split()) >= 12] # Lowered from 15 to 12 words
        random.shuffle(essay_candidates)
        
        current_essays = 0
        for sent in essay_candidates:
            if current_essays >= num_essay:
                break
                
            # Avoid using same sentences as MCQs if possible, but relax this if we don't have enough candidates
            is_duplicate = any(sent in q.get('question', '') or sent in q.get('context', '') for q in quiz_questions)
            if is_duplicate and len(essay_candidates) > num_essay * 2:
                continue
                
            essay = generate_essay(sent)
            if essay:
                quiz_questions.append(essay)
                current_essays += 1
        
        # Phase 3: DESPERATION FALLBACK (if we still don't have enough essays)
        if current_essays < num_essay:
            print(f"Desperation fallback for essays: only found {current_essays}/{num_essay}")
            for sent in sentences: # Try any sentence now
                if current_essays >= num_essay: break
                if any(q.get('type') == 'essay' and q.get('context') == sent for q in quiz_questions): continue
                
                essay = generate_essay(sent)
                if essay:
                    quiz_questions.append(essay)
                    current_essays += 1

        gc.collect() # Final cleanup
        return {"quiz": {"questions": quiz_questions}}
    except Exception as e:
        gc.collect() # Cleanup on error
        print(f"Quiz Generation Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/explain-incorrect")
async def explain_incorrect(data: dict):
    try:
        question = data.get("question", "")
        user_answer = data.get("user_answer", "")
        correct_answer = data.get("correct_answer", "")
        context = data.get("context", "")
        material_text = data.get("material_text", "")
        
        # 1. Clean the question (remove placeholders for better POS tagging if needed)
        clean_question = question.replace('________', correct_answer)
        
        # 2. Extract the main topic - prioritize the CORRECT ANSWER itself
        # This ensures we never get "________" as a topic
        topic = correct_answer
        
        # Optional: if correct answer is too short, try to find a noun in the question
        if len(topic) <= 3:
            tokens = word_tokenize(clean_question)
            tagged = pos_tag(tokens)
            nouns = [word for word, pos in tagged if pos.startswith('N') and len(word) > 4]
            topic = nouns[0] if nouns else correct_answer

        source_excerpt = get_relevant_excerpt(material_text, question, correct_answer, user_answer, context)
        source_reference = source_excerpt or context or clean_question

        safe_topic = re.sub(r'[^a-zA-Z0-9\s-]', '', topic).strip()
        search_query = quote_plus(
            " ".join(filter(None, [
                safe_topic or correct_answer,
                question.replace("________", ""),
                source_reference[:180]
            ]))
        )

        links = [
            {
                "title": f"YouTube: {topic} explained",
                "url": f"https://www.youtube.com/results?search_query={search_query}+explained",
                "type": "video"
            },
            {
                "title": f"Khan Academy: {topic}",
                "url": f"https://www.khanacademy.org/search?page_search_query={search_query}",
                "type": "article"
            },
            {
                "title": f"Google Search: {topic} definition",
                "url": f"https://www.google.com/search?q={search_query}+definition",
                "type": "article"
            }
        ]

        deterministic_payload = build_deterministic_explanation(
            question=question,
            user_answer=user_answer,
            correct_answer=correct_answer,
            context=context,
            source_reference=source_reference,
            topic=topic
        )

        if groq_available:
            try:
                
                prompt = f"""
You are an expert tutor explaining a wrong answer to a university student.
Use the deterministic teaching notes below as the factual base. Expand them, do not replace them with vague filler.

Question:
{question}

Student answer:
{user_answer}

Correct answer:
{correct_answer}

Context from lesson / source material:
{context}

Full material excerpt selected from the original course text:
{source_reference}

Deterministic teaching notes:
{json.dumps(deterministic_payload, ensure_ascii=False, indent=2)}

Task:
1. Explain the concept deeply and clearly in 2-4 short paragraphs.
2. Explain why the student's answer is wrong.
3. Explain the real meaning of the correct answer in this context.
4. Break the concept into a simple concept breakdown with concrete wording.
5. Give one simple memory tip or analogy.
6. Recommend useful search terms for further study.
7. Use the course excerpt directly. Do NOT answer with vague filler like "key term", "important term", or "essential concept".

Return valid JSON only with this schema:
{{
  "explanation": "full explanation with line breaks",
  "why_wrong": "why the selected answer is incorrect",
  "real_meaning": "what the correct answer actually means in this context",
  "concept_breakdown": "a simple breakdown of the idea in student-friendly language",
  "memory_tip": "a short memory aid",
  "suggestions": "what the student should study next",
  "search_terms": ["term 1", "term 2", "term 3"],
  "source_excerpt": "the most relevant part of the course text used for the answer"
}}
"""
                response = model.generate_content(
                    prompt,
                    generation_config={"response_mime_type": "application/json"}
                )
                payload = json.loads(strip_json_markdown(response.text))
                payload = {**deterministic_payload, **payload}
                payload["links"] = links
                if source_reference and not payload.get("source_excerpt"):
                    payload["source_excerpt"] = source_reference[:1200]
                if not payload.get("suggestions"):
                    payload["suggestions"] = deterministic_payload["suggestions"]
                if not payload.get("concept_breakdown"):
                    payload["concept_breakdown"] = deterministic_payload["concept_breakdown"]
                if looks_generic_explanation(payload):
                    payload = {**deterministic_payload, "links": links}
                return payload
            except Exception as gemini_err:
                print(f"Gemini explanation generation failed: {gemini_err}. Falling back to template explanation.")

        return dict(deterministic_payload, links=links)
    except Exception as e:
        print(f"Error configuring Google Gemini: {e}")

def get_gemini_model(model_name="gemini-2.5-flash"):
    """
    Returns a configured GenerativeModel with a robust fallback chain.
    """
    try:
        return genai.GenerativeModel(model_name)
    except Exception as e:
        print(f"Failed to load primary model {model_name}: {e}. Trying fallback chain...")
        try:
            return genai.GenerativeModel("gemini-2.5-flash")
        except Exception:
            try:
                return genai.GenerativeModel("gemini-2.0-flash")
            except Exception:
                try:
                    return genai.GenerativeModel("gemini-pro-latest")
                except Exception:
                    return genai.GenerativeModel("gemini-1.5-flash")

def strip_json_markdown(text: str) -> str:
    """
    Strips markdown code block wrappers (like ```json ... ```) from JSON response strings.
    """
    text = text.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].startswith("```"):
            lines = lines[:-1]
        text = "\n".join(lines).strip()
    
    start_brace = text.find("{")
    start_bracket = text.find("[")
    
    start_idx = -1
    end_idx = -1
    
    if start_brace != -1 and (start_bracket == -1 or start_brace < start_bracket):
        start_idx = start_brace
        end_idx = text.rfind("}")
    elif start_bracket != -1:
        start_idx = start_bracket
        end_idx = text.rfind("]")
        
    if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
        text = text[start_idx:end_idx+1]
        
    return text

def get_relevant_excerpt(material_text: str, question: str, correct_answer: str, user_answer: str, context: str, max_sentences: int = 6, max_chars: int = 2500) -> str:
    """
    Pick the most relevant excerpt from the full study material so the explanation
    model can reason over the actual source instead of only the question snippet.
    """
    if not material_text or not material_text.strip():
        return ""

    clean_text = re.sub(r'\s+', ' ', material_text).strip()
    sentences = [s.strip() for s in sent_tokenize(clean_text) if len(s.strip()) > 20]
    if not sentences:
        return clean_text[:max_chars]

    keyword_sources = " ".join([question or "", correct_answer or "", user_answer or "", context or ""]).lower()
    keywords = {
        token for token in re.findall(r"[a-zA-Z][a-zA-Z\-']+", keyword_sources)
        if len(token) > 3
    }

    scored_sentences = []
    for sentence in sentences:
        sent_lower = sentence.lower()
        score = 0
        for keyword in keywords:
            if keyword in sent_lower:
                score += 2
        if correct_answer and correct_answer.lower() in sent_lower:
            score += 6
        if context and context[:120].lower() in sent_lower:
            score += 4
        if question and question[:120].lower() in sent_lower:
            score += 3
        scored_sentences.append((score, sentence))

    top_sentences = [sentence for score, sentence in sorted(scored_sentences, key=lambda item: item[0], reverse=True) if score > 0][:max_sentences]
    if not top_sentences:
        top_sentences = sentences[:max_sentences]

    excerpt = " ".join(top_sentences)
    return excerpt[:max_chars]

def build_deterministic_explanation(question: str, user_answer: str, correct_answer: str, context: str, source_reference: str, topic: str) -> dict:
    clean_question = re.sub(r'\s+', ' ', (question or '')).strip()
    completed_question = clean_question.replace('________', correct_answer).replace('______', correct_answer).replace('_____', correct_answer)
    completed_question = completed_question.replace('____', correct_answer)

    if not completed_question:
        completed_question = f"The correct phrase is '{correct_answer}' in the given context."

    if source_reference:
        source_excerpt = source_reference[:1200]
    else:
        source_excerpt = context[:1200] if context else completed_question[:1200]

    completed_phrase = completed_question if completed_question else correct_answer
    if completed_phrase == correct_answer and ' and ' in clean_question.lower():
        completed_phrase = f"{correct_answer} and the rest of the technical phrase from the sentence"

    if 'communications technologies' in clean_question.lower() and 'information' in correct_answer.lower():
        real_meaning = (
            "The completed phrase is 'Information and Communications Technologies' (ICT). "
            "It refers to the infrastructure, devices, software, networks, and systems used to process, store, and transmit digital information."
        )
        concept_breakdown = (
            "The blank is not just any noun; it is the word that completes a fixed technical phrase. "
            "Here, 'Information' works because the phrase becomes the standard computing term ICT."
        )
        why_wrong = (
            f"'{user_answer}' does not complete the technical phrase correctly and does not match the definition of modern computing infrastructure."
        )
        memory_tip = "If a blank completes a well-known technical phrase, check the standard term first before guessing a random noun."
    else:
        real_meaning = (
            f"'{correct_answer}' is the exact word or phrase that fits the definition, sentence logic, and subject being discussed in the material."
        )
        concept_breakdown = (
            f"Think of '{correct_answer}' as the word that makes the sentence technically and conceptually correct, not just a related term."
        )
        why_wrong = (
            f"'{user_answer}' sounds related, but it does not fit the sentence structure or the meaning required by the material."
        )
        memory_tip = "Focus on the exact sentence meaning and the technical phrase, not just a familiar-sounding option."

    explanation = (
        f"The completed statement is:\n{completed_phrase}\n\n"
        f"Why '{user_answer}' is wrong:\n{why_wrong}\n\n"
        f"Real meaning in context:\n{real_meaning}\n\n"
        f"Concept breakdown:\n{concept_breakdown}\n\n"
        f"Course logic:\n{source_excerpt}"
    )

    return {
        "explanation": explanation,
        "why_wrong": why_wrong,
        "real_meaning": real_meaning,
        "concept_breakdown": concept_breakdown,
        "memory_tip": memory_tip,
        "suggestions": f"Review the section of the material that defines {correct_answer} and compare it with the surrounding sentence.",
        "source_excerpt": source_excerpt,
        "search_terms": [correct_answer, topic, clean_question[:60]]
    }

def looks_generic_explanation(payload: dict) -> bool:
    text = " ".join([
        str(payload.get("explanation", "")),
        str(payload.get("real_meaning", "")),
        str(payload.get("concept_breakdown", "")),
        str(payload.get("suggestions", ""))
    ]).lower()
    generic_markers = [
        "key term",
        "essential concept",
        "important concept",
        "understanding this concept is essential",
        "reviewing",
        "will help you master this specific concept"
    ]
    return (
        len(str(payload.get("explanation", "")).strip()) < 120
        or any(marker in text for marker in generic_markers)
        or not payload.get("concept_breakdown")
        or not payload.get("real_meaning")
    )


app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- NLP Objective Question Generator ---

def preprocess_text(text):
    """Clean and tokenize text into sentences"""
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    sentences = sent_tokenize(text)
    return [s for s in sentences if len(s.strip()) > 40]

def is_high_quality(sentence):
    """Filter for sentences that sound like facts or definitions"""
    indicators = [
        ' is ', ' are ', ' refers to ', ' defined as ', ' known as ', 
        ' because ', ' therefore ', ' significant ', ' primary ', 
        ' essentially ', ' consists of ', ' includes ', ' process of '
    ]
    # Length check + indicator check
    if len(sentence.split()) < 10:
        return False
        
    return any(ind in sentence.lower() for ind in indicators)

def extract_keywords(sentence):
    """Extract candidate keywords using POS tagging, prioritizing longer nouns"""
    tokens = word_tokenize(sentence)
    tagged = pos_tag(tokens)
    
    # Priority 1: Proper Nouns (NNP) - usually names of theories, places, things
    # Priority 2: Standard Nouns (NN)
    # Priority 3: Adjectives over 6 chars
    candidates = []
    for word, pos in tagged:
        if len(word) < 4 or not word.isalnum():
            continue
            
        if pos == 'NNP':
            candidates.append((word, pos, 3)) # High priority
        elif pos in ('NN', 'NNS'):
            candidates.append((word, pos, 2))
        elif pos == 'JJ' and len(word) > 6:
            candidates.append((word, pos, 1))
            
    # Sort by priority then by length (longer = more specific)
    candidates.sort(key=lambda x: (x[2], len(x[0])), reverse=True)
    return [(c[0], c[1]) for c in candidates]

def get_wordnet_pos(treebank_tag):
    """Map treebank POS tags to WordNet POS tags"""
    if treebank_tag.startswith('J'):
        return wordnet.ADJ
    elif treebank_tag.startswith('V'):
        return wordnet.VERB
    elif treebank_tag.startswith('N'):
        return wordnet.NOUN
    elif treebank_tag.startswith('R'):
        return wordnet.ADV
    else:
        return wordnet.NOUN

_distractor_cache = {}

def generate_distractors(target_word, pos_tag):
    """Generate distractors using WordNet strategies (hypernyms, hyponyms, coordinates)"""
    cache_key = f"{target_word.lower()}_{pos_tag}"
    if cache_key in _distractor_cache:
        return _distractor_cache[cache_key]
        
    distractors = set()
    wn_pos = get_wordnet_pos(pos_tag)
    
    synsets = wordnet.synsets(target_word, pos=wn_pos)
    if not synsets:
        # Try without POS restriction if nothing found
        synsets = wordnet.synsets(target_word)
        
    if synsets:
        synset = synsets[0]
        
        # 1. Hypernym Strategy (Coordinate terms)
        for hypernym in synset.hypernyms():
            for hyponym in hypernym.hyponyms():
                name = hyponym.lemmas()[0].name().replace('_', ' ')
                if name.lower() != target_word.lower():
                    distractors.add(name)
                    if len(distractors) >= 10: break
            if len(distractors) >= 10: break
            
        # 2. Hyponym Strategy
        if len(distractors) < 5:
            for hyponym in synset.hyponyms():
                name = hyponym.lemmas()[0].name().replace('_', ' ')
                if name.lower() != target_word.lower():
                    distractors.add(name)
                    if len(distractors) >= 15: break

    result = list(distractors)
    _distractor_cache[cache_key] = result
    # Keep cache small but reasonable
    if len(_distractor_cache) > 1000:
        _distractor_cache.clear()
        
    return result

def filter_distractors(distractors, target_word, sentence, pos_tag):
    """Apply filtering rules for high-quality distractors"""
    filtered = []
    target_len = len(target_word)
    
    # Pre-tokenize sentence for POS consistency check
    # (Simple check: distractors should likely be valid in context)
    
    for d in distractors:
        d = d.lower()
        # Rule: Not a synonym or duplicate
        if d == target_word.lower() or d in filtered:
            continue
            
        # Rule: Not a substring of the answer
        if d in target_word.lower() or target_word.lower() in d:
            continue
            
        # Rule: Similar POS (approximated by WordNet match or word structure)
        # Rule: Reasonable length (not wildly different from target)
        if abs(len(d) - target_len) > 8:
            continue
            
        filtered.append(d)
        if len(filtered) >= 3:
            break
            
    # Fallback to similar words if WordNet fails
    fallbacks = ["process", "system", "theory", "method", "concept", "approach", "framework"]
    while len(filtered) < 3:
        f = random.choice(fallbacks)
        if f not in filtered and f != target_word.lower():
            filtered.append(f)
            
    return filtered

def score_question(question_text, answer, distractors):
    """Simple scoring module (0-100)"""
    score = 100
    
    # 1. POS consistency (simple check: if answer is capitalized, distractors should be too - ignored for now)
    
    # 2. Option length similarity
    avg_len = sum(len(o) for o in distractors + [answer]) / 4
    for o in distractors + [answer]:
        if abs(len(o) - avg_len) > 10:
            score -= 10
            
    # 3. Duplicate check
    if len(set([answer] + distractors)) < 4:
        score -= 50
        
    return score

def generate_mcq(sentence, target_keyword_info):
    """Build a complete MCQ object following the pipeline"""
    word, pos = target_keyword_info
    
    # 1. Distractors
    raw_distractors = generate_distractors(word, pos)
    filtered = filter_distractors(raw_distractors, word, sentence, pos)
    
    # 2. Scoring
    q_score = score_question(sentence, word, filtered)
    if q_score < 70:
        return None # Reject low quality
    
    # 3. Format
    options = filtered + [word]
    random.shuffle(options)
    
    # Assign A, B, C, D
    mapping = {0: 'A', 1: 'B', 2: 'C', 3: 'D'}
    answer_letter = ''
    formatted_options = []
    for i, opt in enumerate(options):
        letter = mapping[i]
        formatted_options.append(f"{letter}. {opt}")
        if opt == word:
            answer_letter = letter
            
    return {
        "type": "mcq",
        "question": sentence.replace(word, "________"),
        "options": options,
        "formatted_options": formatted_options,
        "answer": word,
        "correct_letter": answer_letter,
        "explanation": f"The term '{word}' correctly completes the context of this statement found in your study material.",
        "score": q_score
    }

def generate_essay(sentence):
    """Generate an essay/short answer question from a sentence"""
    keywords = extract_keywords(sentence)
    if not keywords:
        return None
    
    # List of generic words to NEVER use as essay topics
    blacklist = {
        'term', 'context', 'material', 'section', 'point', 'study', 'question', 
        'text', 'sentence', 'subject', 'relationship', 'importance', 'overall',
        'significance', 'discussion', 'concept', 'structure', 'element', 'core'
    }
    
    # 1. Focus on NOUNS (NN, NNP, NNS)
    # 2. Filter out blacklist words
    # 3. Require minimum length of 6 for better quality
    candidates = [word for word, pos in keywords 
                 if pos.startswith('N') 
                 and word.lower() not in blacklist 
                 and len(word) >= 6]
    
    if not candidates:
        # Fallback to any noun over 5 chars if no specific candidate found
        candidates = [word for word, pos in keywords if pos.startswith('N') and len(word) > 5]
        
    if not candidates:
        return None
    
    # Pick the LONGEST noun - it is usually the most specific technical term
    subject = max(candidates, key=len)
    
    prompts = [
        f"Explain the significance of '{subject}' based on your study material.",
        f"How does '{subject}' relate to the key themes discussed in this section?",
        f"Discuss the role and impact of '{subject}' as described in the provided text.",
        f"Looking at your material, summarize the most important points regarding '{subject}'."
    ]
    
    return {
        "type": "essay",
        "question": random.choice(prompts),
        "context": sentence,
        "answer": sentence, # Context serves as the reference answer
        "explanation": f"Analyze your response above. Does it correctly address the role of '{subject}'? The reference material says: \"{sentence}\""
    }

# --- API Endpoints ---

@app.get("/")
async def root():
    return {"status": "OK", "service": "TutorBuddy AI Core", "version": "2.0 (NLP Pipeline)"}

@app.get("/health")
async def health():
    return {"status": "OK"}

@app.post("/extract-text")
async def extract_text(file: UploadFile = File(...)):
    filename = file.filename.lower()
    print(f"--- Ultra-Lean Extraction: {filename} ---")
    
    # Use a temporary file to keep RAM free
    with tempfile.NamedTemporaryFile(delete=False) as tmp:
        try:
            shutil.copyfileobj(file.file, tmp)
            tmp_path = tmp.name
            tmp.close() # Close so other libs can open it

            if filename.endswith('.pdf'):
                reader = PyPDF2.PdfReader(tmp_path)
                pages_content = []
                # Extreme safety: 30 pages / 30k chars
                page_limit = min(len(reader.pages), 35)
                
                total_chars = 0
                for i in range(page_limit):
                    page_text = reader.pages[i].extract_text() or ""
                    pages_content.append(page_text)
                    total_chars += len(page_text)
                    if total_chars > 30000: break
                
                return {"text": "".join(pages_content)[:30000]}
                
            elif filename.endswith('.docx'):
                import docx
                doc = docx.Document(tmp_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                return {"text": text[:30000]}
                
            elif filename.endswith('.txt'):
                with open(tmp_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return {"text": f.read(30000)}
            else:
                raise HTTPException(status_code=400, detail="Unsupported format")
        
        except Exception as e:
            print(f"Lean Extraction Error: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Failed: {str(e)}")
        finally:
            # Cleanup!
            if 'tmp_path' in locals() and os.path.exists(tmp_path):
                os.remove(tmp_path)
            # Force RAM cleanup
            gc.collect()
            await file.close()
 
@app.post("/generate-quiz")
async def generate_quiz(data: dict):
    gc.collect() # Clean up before starting
    try:
        text = data.get("text", "")
        num_requested = data.get("num_questions", 35) # Default total 35
        num_mcq = data.get("num_mcq", 30)
        num_essay = data.get("num_essay", 5)
        
        print(f"--- Quiz Generation Request Received ---")
        print(f"Text length: {len(text)} characters. Targets: {num_mcq} MCQ, {num_essay} Essay")
        
        if not text:
            raise HTTPException(status_code=400, detail="No text provided")
 
        # Lowered limit for better stability on free tier
        if len(text) > 30000:
            print(f"Truncating text from {len(text)} to 30,000 chars for performance.")
            text = text[:30000]

        # --- Premium Google Gemini Path ---
        if groq_available:
            print("Running Premium Groq Quiz Generator...")
            try:
                

                mcq_instruction = f"The quiz MUST contain exactly {num_mcq} Multiple Choice Questions (MCQ)." if num_mcq > 0 else ""
                essay_instruction = f"The quiz MUST contain exactly {num_essay} Essay/Short Answer questions." if num_essay > 0 else ""
                
                mcq_format = """
                For MCQs:
                - The questions must be deep, conceptual, and check real-world logic or application.
                - Avoid simple blank fill-ins or direct word matching.
                - Each MCQ must have exactly 4 options.
                - The "options" list must contain the plain texts.
                - The "formatted_options" list must be prefixed with "A. ", "B. ", "C. ", "D. ".
                - One option must be the correct answer.
                - "correct_letter" must be 'A', 'B', 'C', or 'D', matching the correct option.
                - Include a detailed educational "explanation" for why it is correct.
                - A quality "score" between 80 and 100.
                """ if num_mcq > 0 else ""
                
                essay_format = """
                For Essays:
                - The questions must ask the student to explain core concepts, relationships between mechanisms, or summary of main themes in the text.
                - "context" must be a supporting reference sentence from the text.
                - "answer" must be a comprehensive reference master answer.
                - "explanation" must be educational criteria for a correct response.
                """ if num_essay > 0 else ""
                
                prompt = f"""
                Generate a study quiz based on the text below. 
                {mcq_instruction}
                {essay_instruction}
                
                {mcq_format}
                
                {essay_format}
                
                Output the result STRICTLY as a JSON object matching this schema:
                {{
                  "questions": [

                     {{
                       "type": "mcq",
                       "question": "question text",
                       "options": ["option 1", "option 2", "option 3", "option 4"],
                       "formatted_options": ["A. option 1", "B. option 2", "C. option 3", "D. option 4"],
                       "answer": "option text matching the correct answer",
                       "correct_letter": "A",
                       "explanation": "detailed educational explanation of why it is correct",
                       "score": 95
                     }},
                     ...
                     {{
                       "type": "essay",
                       "question": "conceptual question about core themes",
                       "context": "supporting context from the text",
                       "answer": "comprehensive reference master answer",
                       "explanation": "educational criteria for a correct response"
                     }}
                  ]
                }}
                
                CRITICAL INSTRUCTION: DO NOT summarize the text. DO NOT write any introductory or concluding text. 
                You must ONLY output the raw JSON object containing the questions array. No markdown, no explanations outside the JSON.
                
                Source material text:
                {text}
                """
                
                text_content = get_groq_completion(prompt, response_format={"type": "json_object"})
                text_content = strip_json_markdown(text_content)
                quiz_data = json.loads(text_content)
                if "questions" in quiz_data and len(quiz_data["questions"]) > 0:
                    print(f"Groq Quiz Generation successful! Generated {len(quiz_data['questions'])} questions.")
                    return {"quiz": quiz_data}
                else:
                    print("Groq response did not contain questions. Falling back to local NLTK...")
            except Exception as e:
                print(f"Groq quiz generation failed ({e}). Falling back to local NLTK...")

        # --- Local NLTK Fallback Path ---
        print("Running Local NLTK Quiz Generator...")
        sentences = preprocess_text(text)
        if not sentences:
            raise HTTPException(status_code=400, detail="Text too short or invalid to generate questions")
            
        random.shuffle(sentences)
        
        quiz_questions = []
        used_words = set()
        
        # Phase 1: MCQs
        mcq_count = 0
        for i, sent in enumerate(sentences):
            if mcq_count >= num_mcq:
                break
                
            # NEW: prioritize "Educational" sentences
            if not is_high_quality(sent) and len(sentences) > num_mcq * 3:
                continue

            if i % 10 == 0:
                print(f"Processing MCQ sentence {i}/{len(sentences)}... Found {mcq_count} so far.")

            keywords = extract_keywords(sent)
            if not keywords or len(keywords) < 2: 
                continue
                
            # Pick a lucky keyword
            random.shuffle(keywords)
            for kw_info in keywords:
                word = kw_info[0]
                if word.lower() in used_words: 
                    continue
                    
                mcq = generate_mcq(sent, kw_info)
                if mcq:
                    quiz_questions.append(mcq)
                    used_words.add(word.lower())
                    mcq_count += 1
                    break
        
        # Phase 1.5: MCQ Fallback (if we still don't have enough MCQs)
        current_mcq = len([q for q in quiz_questions if q['type'] == 'mcq'])
        if current_mcq < num_mcq:
            print(f"MCQ Fallback: only found {current_mcq}/{num_mcq}")
            for sent in sentences:
                if len([q for q in quiz_questions if q['type'] == 'mcq']) >= num_mcq: break
                # Simplified MCQ for fallback: just pick a word
                words = [w for w in word_tokenize(sent) if len(w) > 5 and w.isalnum()]
                for w in words:
                    if w.lower() in used_words: continue
                    mcq = generate_mcq(sent, (w, 'NN'))
                    if mcq:
                        quiz_questions.append(mcq)
                        used_words.add(w.lower())
                        break
        
        # Phase 2: Essays
        essay_candidates = [s for s in sentences if len(s.split()) >= 12] # Lowered from 15 to 12 words
        random.shuffle(essay_candidates)
        
        current_essays = 0
        for sent in essay_candidates:
            if current_essays >= num_essay:
                break
                
            # Avoid using same sentences as MCQs if possible, but relax this if we don't have enough candidates
            is_duplicate = any(sent in q.get('question', '') or sent in q.get('context', '') for q in quiz_questions)
            if is_duplicate and len(essay_candidates) > num_essay * 2:
                continue
                
            essay = generate_essay(sent)
            if essay:
                quiz_questions.append(essay)
                current_essays += 1
        
        # Phase 3: DESPERATION FALLBACK (if we still don't have enough essays)
        if current_essays < num_essay:
            print(f"Desperation fallback for essays: only found {current_essays}/{num_essay}")
            for sent in sentences: # Try any sentence now
                if current_essays >= num_essay: break
                if any(q.get('type') == 'essay' and q.get('context') == sent for q in quiz_questions): continue
                
                essay = generate_essay(sent)
                if essay:
                    quiz_questions.append(essay)
                    current_essays += 1

        gc.collect() # Final cleanup
        return {"quiz": {"questions": quiz_questions}}
    except Exception as e:
        gc.collect() # Cleanup on error
        print(f"Quiz Generation Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/explain-incorrect")
async def explain_incorrect(data: dict):
    try:
        question = data.get("question", "")
        user_answer = data.get("user_answer", "")
        correct_answer = data.get("correct_answer", "")
        context = data.get("context", "")
        material_text = data.get("material_text", "")
        
        # 1. Clean the question (remove placeholders for better POS tagging if needed)
        clean_question = question.replace('________', correct_answer)
        
        # 2. Extract the main topic - prioritize the CORRECT ANSWER itself
        # This ensures we never get "________" as a topic
        topic = correct_answer
        
        # Optional: if correct answer is too short, try to find a noun in the question
        if len(topic) <= 3:
            tokens = word_tokenize(clean_question)
            tagged = pos_tag(tokens)
            nouns = [word for word, pos in tagged if pos.startswith('N') and len(word) > 4]
            topic = nouns[0] if nouns else correct_answer

        source_excerpt = get_relevant_excerpt(material_text, question, correct_answer, user_answer, context)
        source_reference = source_excerpt or context or clean_question

        safe_topic = re.sub(r'[^a-zA-Z0-9\s-]', '', topic).strip()
        search_query = quote_plus(
            " ".join(filter(None, [
                safe_topic or correct_answer,
                question.replace("________", ""),
                source_reference[:180]
            ]))
        )

        links = [
            {
                "title": f"YouTube: {topic} explained",
                "url": f"https://www.youtube.com/results?search_query={search_query}+explained",
                "type": "video"
            },
            {
                "title": f"Khan Academy: {topic}",
                "url": f"https://www.khanacademy.org/search?page_search_query={search_query}",
                "type": "article"
            },
            {
                "title": f"Google Search: {topic} definition",
                "url": f"https://www.google.com/search?q={search_query}+definition",
                "type": "article"
            }
        ]

        deterministic_payload = build_deterministic_explanation(
            question=question,
            user_answer=user_answer,
            correct_answer=correct_answer,
            context=context,
            source_reference=source_reference,
            topic=topic
        )

        if groq_available:
            try:
                
                prompt = f"""
You are an expert tutor explaining a wrong answer to a university student.
Use the deterministic teaching notes below as the factual base. Expand them, do not replace them with vague filler.

Question:
{question}

Student answer:
{user_answer}

Correct answer:
{correct_answer}

Context from lesson / source material:
{context}

Full material excerpt selected from the original course text:
{source_reference}

Deterministic teaching notes:
{json.dumps(deterministic_payload, ensure_ascii=False, indent=2)}

Task:
1. Explain the concept deeply and clearly in 2-4 short paragraphs.
2. Explain why the student's answer is wrong.
3. Explain the real meaning of the correct answer in this context.
4. Break the concept into a simple concept breakdown with concrete wording.
5. Give one simple memory tip or analogy.
6. Recommend useful search terms for further study.
7. Use the course excerpt directly. Do NOT answer with vague filler like "key term", "important term", or "essential concept".

Return valid JSON only with this schema:
{{
  "explanation": "full explanation with line breaks",
  "why_wrong": "why the selected answer is incorrect",
  "real_meaning": "what the correct answer actually means in this context",
  "concept_breakdown": "a simple breakdown of the idea in student-friendly language",
  "memory_tip": "a short memory aid",
  "suggestions": "what the student should study next",
  "search_terms": ["term 1", "term 2", "term 3"],
  "source_excerpt": "the most relevant part of the course text used for the answer"
}}
"""
                response = model.generate_content(
                    prompt,
                    generation_config={"response_mime_type": "application/json"}
                )
                payload = json.loads(strip_json_markdown(response.text))
                payload = {**deterministic_payload, **payload}
                payload["links"] = links
                if source_reference and not payload.get("source_excerpt"):
                    payload["source_excerpt"] = source_reference[:1200]
                if not payload.get("suggestions"):
                    payload["suggestions"] = deterministic_payload["suggestions"]
                if not payload.get("concept_breakdown"):
                    payload["concept_breakdown"] = deterministic_payload["concept_breakdown"]
                if looks_generic_explanation(payload):
                    payload = {**deterministic_payload, "links": links}
                return payload
            except Exception as gemini_err:
                print(f"Gemini explanation generation failed: {gemini_err}. Falling back to template explanation.")

        explanation = (
            f"You selected '{user_answer}', but the correct answer is '{correct_answer}'.\n\n"
            f"In this question, '{topic}' is not just a word to memorize — it refers to the idea that best fits the sentence and the subject being discussed.\n\n"
            f"Why it matters: the question is testing whether you can identify the term that makes the statement logically and conceptually correct. "
            f"'{user_answer}' sounds related, but it does not match the meaning required by the sentence.\n\n"
            f"From the course material, the most relevant section is:\n\n"
            f"{source_reference}\n\n"
            f"To understand it properly, connect the term to that exact material and ask: what concept best fits the age, qualification, or context being described?"
        )

        return {
            "explanation": explanation,
            "why_wrong": f"'{user_answer}' does not fit the meaning required by the sentence.",
            "real_meaning": f"'{correct_answer}' is the concept that best matches the context of the question.",
            "concept_breakdown": f"Think of '{correct_answer}' as the exact idea the sentence is asking for, not just a related-sounding word.",
            "memory_tip": f"Link the word to the sentence logic, not just to a familiar-sounding term.",
            "suggestions": f"Review {topic} in the original course material and compare it with nearby keywords in the text.",
            "source_excerpt": source_reference[:1200] if source_reference else "",
            "links": links
        }
    except Exception as e:
        print(f"Explanation Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/adaptive-feedback")
async def adaptive_feedback(data: dict):
    try:
        score = int(data.get("score", 0) or 0)
        difficulty = str(data.get("difficulty", "medium")).lower()
        performance_score = data.get("performance_score")
        course_title = data.get("course_title", "") or "this course"
        quiz_title = data.get("quiz_title", "") or "this quiz"

        if score >= 75:
            level = "hard"
            headline = "You’ve earned a harder challenge"
            why = f"Your score of {score}% shows strong control of the material, so the next quiz will move into harder reasoning and application."
            next_focus = "Expect deeper concept checks, trickier distractors, and scenario-based questions."
            encouragement = "Keep pushing — you’re ready for more demanding questions."
        elif score < 50:
            level = "easy"
            headline = "Time for a revision-focused quiz"
            why = f"Your score of {score}% suggests this topic still needs reinforcement, so the next quiz will focus on simpler revision and core recall."
            next_focus = "Expect clearer wording, more direct prompts, and questions that revisit the basics."
            encouragement = "This is how progress works — tighten the basics first, then level up."
        else:
            level = "medium"
            headline = "Your next quiz will stay balanced"
            why = f"Your score of {score}% shows you’re in the middle zone, so the next quiz will keep a balanced mix of recall and understanding."
            next_focus = "Expect a mix of straightforward questions and some concept application."
            encouragement = "You’re progressing steadily — keep building consistency."

        base_payload = {
            "level": level,
            "headline": headline,
            "message": why,
            "why_this_level": why,
            "next_focus": next_focus,
            "encouragement": encouragement,
            "course_title": course_title,
            "quiz_title": quiz_title,
            "score": score,
            "performance_score": performance_score,
            "difficulty": difficulty
        }

        if groq_available:
            try:
                
                prompt = f"""
You are a warm academic tutor explaining why the student's next quiz level changed.

Student score: {score}%
Current adaptive level: {difficulty}
Recent performance score: {performance_score}
Course title: {course_title}
Quiz title: {quiz_title}

Write a short, encouraging explanation that tells the student:
1. why this level was chosen,
2. what kind of questions they should expect next,
3. how this helps their learning.

Use a professional but supportive tone. Avoid vague filler.

Return valid JSON only with this schema:
{{
  "level": "easy | medium | hard",
  "headline": "short headline",
  "message": "main explanation in 2-4 short sentences",
  "why_this_level": "why the level was chosen",
  "next_focus": "what the student should expect next",
  "encouragement": "short encouraging line"
}}
"""
                response = model.generate_content(
                    prompt,
                    generation_config={"response_mime_type": "application/json"}
                )
                payload = json.loads(strip_json_markdown(response.text))
                payload = {**base_payload, **payload}
                return payload
            except Exception as gemini_err:
                print(f"Adaptive feedback Gemini generation failed: {gemini_err}. Falling back to deterministic message.")

        return base_payload
    except Exception as e:
        print(f"Adaptive feedback error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/analyze-weakness")
async def analyze_weakness(data: dict):
    try:
        incorrect_data = data.get("incorrect_data", [])
        if not incorrect_data:
            return {
                "weaknesses": [],
                "recommendations": (
                    "You do not have enough mistake history yet. Start with one short quiz, review the course excerpt, "
                    "and write down any repeated terms or definitions you missed."
                )
            }
            
        all_text = " ".join([d.get("question", "") + " " + d.get("correct_answer", "") for d in incorrect_data])
        
        # Use existing POS logic for consistency
        tokens = word_tokenize(all_text)
        tagged = pos_tag(tokens)
        generic_terms = {
            'question', 'questions', 'answer', 'answers', 'unknown', 'n/a', 'review',
            'concept', 'concepts', 'study', 'material', 'materials', 'topic', 'topics'
        }
        keywords = [
            word for word, pos in tagged
            if pos in ('NN', 'NNS', 'NNP')
            and len(word) > 4
            and word.lower() not in generic_terms
        ]
            
        common = Counter(keywords).most_common(3)
        weaknesses = [k[0] for k in common]
        
        return {
            "weaknesses": weaknesses,
            "recommendations": (
                f"Focus on reviewing concepts related to: {', '.join(weaknesses)}."
                if weaknesses
                else "No strong weak spot was detected yet. Try another quiz and revisit the latest lesson excerpt for repeated key terms."
            )
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/evaluate-essay")
async def evaluate_essay(data: dict):
    try:
        question = data.get("question", "")
        student_answer = data.get("student_answer", "")
        context = data.get("context", "") # This is the source sentence
        
        # 1. Extract the topic from the question (using the quotes we added earlier)
        topic_match = re.search(r"'(.*?)'", question)
        topic = topic_match.group(1) if topic_match else "this concept"
        
        if not student_answer or len(student_answer) < 5:
            return {
                "score": 0,
                "relevance": "None",
                "feedback": "Your answer is too short to be evaluated. Please provide a more detailed explanation.",
                "improvement_tips": "A good essay should be at least 2-3 sentences long and use specific terminology.",
                "master_interpretation": f"Based on the material, '{topic}' is described as: \"{context}\"",
                "links": []
            }

        # 2. Score Calculation (Improved overlap logic)
        context_words = set(word_tokenize(context.lower()))
        student_words = set(word_tokenize(student_answer.lower()))
        overlap = student_words.intersection(context_words)
        
        stop_words = set(['the', 'a', 'is', 'are', 'in', 'at', 'on', 'of', 'and', 'for', 'with', 'this', 'that', 'to', 'it', 'from', 'an'])
        meaningful_overlap = [w for w in overlap if len(w) > 3 and w not in stop_words]
        
        score = min(len(meaningful_overlap) * 25, 100) 
        relevance = "High" if score >= 75 else "Medium" if score >= 40 else "Low"
        
        # 3. Dynamic Feedback
        feedback = f"AI Evaluation for '{topic}': "
        if meaningful_overlap:
            feedback += f"You correctly mentioned: {', '.join(meaningful_overlap[:3])}. "
        
        if score < 50:
            tips = "Try to be more specific. Your answer is a bit general. Include more details from the text."
        else:
            tips = "Great job! To get 100%, ensure you explain the relationship between this and other core concepts."

        # 4. Generate Resources for this specific essay topic
        safe_query = re.sub(r'[^a-zA-Z0-9\s]', '', topic).strip().replace(' ', '+')
        links = [
            { "title": f"Mastering {topic} (Video Guide)", "url": f"https://www.youtube.com/results?search_query={safe_query}+explained", "type": "video" },
            { "title": f"Deep Dive: {topic} Resources", "url": f"https://www.khanacademy.org/search?page_search_query={safe_query}", "type": "article" }
        ]

        return {
            "score": score,
            "relevance": relevance,
            "feedback": feedback,
            "improvement_tips": tips,
            "master_concept": f"To master this, remember that the material defines {topic} as follows: {context}",
            "links": links
        }
    except Exception as e:
        print(f"Essay Evaluation Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/summarize")
async def summarize(data: dict):
    try:
        text = data.get("text", "")
        if not text:
            raise HTTPException(status_code=400, detail="No text provided")
            
        if groq_available:
            print("Generating Premium Groq Summary...")
            try:
                
                prompt = f"""
                Create a structured, highly educational, and comprehensive summary of the study notes/material provided below.
                The summary MUST be beautifully formatted in GitHub-flavored Markdown.
                Include the following sections:
                - 💡 **Core Concepts**: Define the most important vocabulary, theories, and concepts.
                - 📌 **Key Takeaways**: Bullet points highlighting the main arguments, mechanisms, or rules.
                - 📖 **Detailed Summary**: A clear narrative explaining the text in detail, breaking it down into logical subsections.
                
                Keep the tone encouraging, academic, and engaging.
                
                Source text to summarize:
                {text}
                """
                text_content = get_groq_completion(prompt)
                return {"summary": text_content}
            except Exception as e:
                print(f"Groq summary generation failed: {e}. Falling back to local summarization.")
                
        # --- Local Fallback Summarization ---
        print("Generating Local NLTK Summary...")
        # clean extra spaces and split sentences
        clean_text = re.sub(r'\s+', ' ', text).strip()
        sentences = sent_tokenize(clean_text)
        if not sentences:
            return {"summary": "### Empty Material\n\nNo text content was found in this study material to summarize."}
            
        words = word_tokenize(clean_text.lower())
        stop_words = set(['the', 'a', 'is', 'are', 'in', 'at', 'on', 'of', 'and', 'for', 'with', 'this', 'that', 'to', 'it', 'from', 'an', 'by', 'as', 'was', 'were', 'or'])
        freq_dict = {}
        for w in words:
            if w.isalnum() and w not in stop_words:
                freq_dict[w] = freq_dict.get(w, 0) + 1
                
        sentence_scores = {}
        for sent in sentences:
            if len(sent.split()) < 8: continue
            for word in word_tokenize(sent.lower()):
                if word in freq_dict:
                    sentence_scores[sent] = sentence_scores.get(sent, 0) + freq_dict[word]
                    
        top_sentences = sorted(sentence_scores, key=sentence_scores.get, reverse=True)[:5]
        summary_text = "\n\n".join(top_sentences)
        
        markdown_summary = "### 💡 Core Concepts & Summary (Offline Fallback)\n\n"
        markdown_summary += "Below are the key sentences extracted from your material using text frequency analysis:\n\n"
        for i, sent in enumerate(top_sentences):
            markdown_summary += f" - {sent}\n"
            
        return {"summary": markdown_summary}
    except Exception as e:
        print(f"Summarization Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/ocr-evaluate")
async def ocr_evaluate(file: UploadFile = File(...)):
    """
    Stand-alone OCR transcription and educational analysis endpoint.
    Transcribes handwritten or printed notes and estimates word count, readability, and feedback.
    """
    try:
        filename = file.filename.lower()
        print(f"--- General OCR Scan: {filename} ---")
        img_bytes = await file.read()
        
        # --- Premium Gemini Vision Path ---
        if groq_available:
            print("Processing general OCR via Gemini Vision...")
            try:
                
                image_part = {
                    "mime_type": file.content_type or "image/png",
                    "data": img_bytes
                }
                prompt = """
                Analyze this image containing handwritten or printed study notes.
                
                Please perform the following actions:
                1. Transcribe all readable text from the image as accurately as possible into "text". Keep formatting, bullet points, and structure intact.
                2. Estimate the "word_count" of the transcribed text.
                3. Classify "readability" as one of: "Highly Legible", "Clear & Legible", "Moderately Clear", or "Hard to Read".
                4. Provide encouraging, supportive educational "feedback" summarizing the notes, main topics covered, and suggesting next study steps.
                
                Output the evaluation strictly as a JSON object matching this schema:
                {
                  "text": "the exact transcribed text...",
                  "evaluation": {
                    "word_count": 120,
                    "readability": "Clear & Legible",
                    "feedback": "..."
                  }
                }
                """
                try:
                    response = model.generate_content(
                        [image_part, prompt],
                        generation_config={"response_mime_type": "application/json"}
                    )
                except Exception as first_err:
                    print(f"General OCR Gemini vision failed ({first_err}), trying fallback model...")
                    fallback_model = get_gemini_model("gemini-2.0-flash")
                    response = fallback_model.generate_content(
                        [image_part, prompt],
                        generation_config={"response_mime_type": "application/json"}
                    )
                
                text_content = strip_json_markdown(response.text)
                result = json.loads(text_content)
                print(f"General OCR Gemini transcription complete! Word count: {result.get('evaluation', {}).get('word_count')}")
                return result
            except Exception as gemini_err:
                print(f"General OCR Gemini vision failed: {gemini_err}. Falling back to Tesseract...")

        # --- Local Tesseract Fallback Path ---
        print("Processing general OCR via Local Tesseract...")
        from PIL import Image, ImageOps, ImageFilter
        import io
        
        try:
            image = Image.open(io.BytesIO(img_bytes))
            image = ImageOps.exif_transpose(image)
            image = image.convert("L")
            image = ImageOps.autocontrast(image)

            if image.width < 1600:
                new_width = image.width * 2
                new_height = image.height * 2
                image = image.resize((new_width, new_height))

            extracted_text = pytesseract.image_to_string(image, config="--oem 3 --psm 6")

            if not extracted_text.strip():
                sharpened_image = image.filter(ImageFilter.SHARPEN)
                extracted_text = pytesseract.image_to_string(sharpened_image, config="--oem 3 --psm 11")
        except Exception as ocr_err:
            print(f"Local Tesseract OCR failed: {ocr_err}")
            extracted_text = ""
            
        if not extracted_text or len(extracted_text.strip()) < 5:
            return {
                "text": "",
                "evaluation": {
                    "word_count": 0,
                    "readability": "Unreadable",
                    "feedback": "Your note image was processed, but no valid text could be extracted. Please make sure the image is clear, upright, and well-lit."
                }
            }
            
        word_count = len(extracted_text.split())
        return {
            "text": extracted_text,
            "evaluation": {
                "word_count": word_count,
                "readability": "Legible (Local OCR)",
                "feedback": f"Successfully extracted {word_count} words of text from your note using local OCR. You can now use these notes to generate quizzes or summary content!"
            }
        }
    except Exception as e:
        print(f"General OCR Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        await file.close()

@app.post("/evaluate-handwritten")
async def evaluate_handwritten(
    file: UploadFile = File(...),
    question: str = File(...),
    reference_answer: str = File(...)
):
    try:
        filename = file.filename.lower()
        print(f"--- Evaluate Handwritten: {filename} ---")
        img_bytes = await file.read()
        
        # --- Premium Gemini Vision Path ---
        if groq_available:
            print("Evaluating Handwritten Answer via Gemini Vision...")
            try:
                
                image_part = {
                    "mime_type": file.content_type or "image/png",
                    "data": img_bytes
                }
                prompt = f"""
                Analyze this handwritten answer to the following academic essay question.
                
                Question: "{question}"
                Reference/Master Answer Context: "{reference_answer}"
                
                Please perform the following actions:
                1. Transcribe the handwritten text from the image as accurately as possible into "extracted_text".
                2. Evaluate the correctness of the answer against the reference answer.
                3. Grade the answer on a scale of 0 to 100 as "score".
                4. Set "relevance" to "High", "Medium", or "Low" based on how well they addressed the core question.
                5. Provide supportive, constructive "feedback" explaining what they got right and what was missing.
                6. Give concrete "improvement_tips" for how they can secure a higher score next time.
                7. Provide a "master_concept" summarizing the ideal understanding.
                
                Output the evaluation strictly as a JSON object matching this schema:
                {{
                  "extracted_text": "the exact transcribed handwriting",
                  "score": 85,
                  "relevance": "High",
                  "feedback": "...",
                  "improvement_tips": "...",
                  "master_concept": "...",
                  "links": [
                     {{"title": "Khan Academy Tutorial", "url": "https://www.khanacademy.org/search?page_search_query=handwritten"}}
                  ]
                }}
                """
                try:
                    response = model.generate_content(
                        [image_part, prompt],
                        generation_config={"response_mime_type": "application/json"}
                    )
                except Exception as first_err:
                    print(f"Handwritten vision evaluation with primary model failed ({first_err}), attempting fallback...")
                    fallback_model = get_gemini_model("gemini-2.0-flash")
                    try:
                        response = fallback_model.generate_content(
                            [image_part, prompt],
                            generation_config={"response_mime_type": "application/json"}
                        )
                    except Exception:
                        response = fallback_model.generate_content([image_part, prompt])
                
                text_content = strip_json_markdown(response.text)
                result = json.loads(text_content)
                print(f"Gemini vision grading complete! Score: {result.get('score')}%")
                return result
            except Exception as gemini_err:
                print(f"Gemini vision grading failed: {gemini_err}. Falling back to Tesseract...")

        # --- Local Tesseract Fallback Path ---
        print("Evaluating Handwritten Answer via Local Tesseract + NLP...")
        from PIL import Image
        import io
        
        try:
            image = Image.open(io.BytesIO(img_bytes))
            extracted_text = pytesseract.image_to_string(image)
        except Exception as ocr_err:
            print(f"Local Tesseract OCR failed: {ocr_err}")
            extracted_text = ""
            
        topic_match = re.search(r"'(.*?)'", question)
        topic = topic_match.group(1) if topic_match else "this concept"
        
        if not extracted_text or len(extracted_text.strip()) < 5:
            return {
                "extracted_text": "",
                "score": 0,
                "relevance": "None",
                "feedback": "Your handwritten answer image was processed, but no valid text could be extracted. Please make sure the image is clear, upright, and well-lit.",
                "improvement_tips": "Try taking the picture again in a well-lit environment and focus the camera on the text.",
                "master_concept": f"Based on the material, '{topic}' is described as: \"{reference_answer}\"",
                "links": []
            }
            
        # Score calculation (overlap logic)
        context_words = set(word_tokenize(reference_answer.lower()))
        student_words = set(word_tokenize(extracted_text.lower()))
        overlap = student_words.intersection(context_words)
        
        stop_words = set(['the', 'a', 'is', 'are', 'in', 'at', 'on', 'of', 'and', 'for', 'with', 'this', 'that', 'to', 'it', 'from', 'an'])
        meaningful_overlap = [w for w in overlap if len(w) > 3 and w not in stop_words]
        
        score = min(len(meaningful_overlap) * 25, 100) 
        relevance = "High" if score >= 75 else "Medium" if score >= 40 else "Low"
        
        feedback = f"Local OCR Evaluation for '{topic}': "
        if meaningful_overlap:
            feedback += f"We detected these matching keywords in your handwriting: {', '.join(meaningful_overlap[:3])}. "
        else:
            feedback += "No clear concept overlap detected in the transcribed text."
            
        if score < 50:
            tips = "Ensure your handwritten answer includes specific key terms and definitions from the study notes."
        else:
            tips = "Excellent handwriting match! You hit the core concepts."
            
        safe_query = re.sub(r'[^a-zA-Z0-9\s]', '', topic).strip().replace(' ', '+')
        links = [
            { "title": f"YouTube Explanation", "url": f"https://www.youtube.com/results?search_query={safe_query}", "type": "video" },
            { "title": f"Google Search", "url": f"https://www.google.com/search?q={safe_query}", "type": "article" }
        ]
        
        return {
            "extracted_text": extracted_text,
            "score": score,
            "relevance": relevance,
            "feedback": feedback,
            "improvement_tips": tips,
            "master_concept": f"The material definition for reference: {reference_answer}",
            "links": links
        }
    except Exception as e:
        print(f"Handwritten Evaluation Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        await file.close()


@app.post("/get-hint")
async def get_hint(data: dict):
    try:
        question = data.get("question", "")
        correct_answer = data.get("correct_answer", "")
        
        if not question:
            raise HTTPException(status_code=400, detail="Missing question")
            
        prompt = f"""
You are a helpful AI tutor. A student is struggling with the following question:
Question: {question}
Correct Answer: {correct_answer}

Provide a brief, encouraging hint that points them in the right direction without directly giving away the answer. Keep it to 1-2 short sentences.
Return valid JSON only:
{{
    "hint": "your hint here"
}}
"""
        chat_completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "You are a helpful study buddy AI. You output strictly valid JSON."},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0.7
        )
        payload = json.loads(chat_completion.choices[0].message.content)
        return payload
    except Exception as e:
        print(f"Hint Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/chat-material")
async def chat_material(data: dict):
    try:
        context_text = data.get("text", "")
        question = data.get("question", "")
        
        if not context_text or not question:
            raise HTTPException(status_code=400, detail="Missing text or question")
            
        # Limit context to avoid hitting token limits for very large PDFs
        max_chars = 30000
        if len(context_text) > max_chars:
            context_text = context_text[:max_chars] + "... [Content Truncated]"

        prompt = f"""
        You are a helpful and expert AI Study Tutor. 
        Your student is asking a question about a specific study material.
        
        STUDY MATERIAL CONTEXT (First {max_chars} characters):
        {context_text}
        
        STUDENT QUESTION:
        {question}
        
        INSTRUCTIONS:
        1. Answer the student's question accurately and clearly.
        2. Base your answer PRIMARILY on the provided STUDY MATERIAL CONTEXT.
        3. If the answer is not in the context, you may use your general knowledge, but mention that it goes beyond the provided text.
        4. Keep your answer concise (1-3 paragraphs) unless more detail is specifically requested.
        5. Format your response in clean Markdown (use bolding, bullet points, etc. where helpful).
        
        Return a strictly valid JSON object with the key "answer" containing your Markdown response.
        Example:
        {{
            "answer": "Here is the explanation... **bold text**."
        }}
        """
        
        if groq_available:
            try:
                chat_completion = groq_client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[
                        {"role": "system", "content": "You are a helpful study buddy AI. You output strictly valid JSON."},
                        {"role": "user", "content": prompt}
                    ],
                    response_format={"type": "json_object"},
                    temperature=0.3
                )
                payload = json.loads(chat_completion.choices[0].message.content)
                return payload
            except Exception as e:
                print(f"Groq Chat Error: {e}")
                raise e
        else:
            raise Exception("Groq API key not configured")
            
    except Exception as e:
        print(f"Chat Material Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
